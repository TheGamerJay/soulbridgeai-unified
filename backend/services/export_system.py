"""
Export System Service
Handle export functionality for poems, stories, and writing content
"""

import logging
import json
import tempfile
import zipfile
from datetime import datetime
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import io

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("ReportLab not available - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX export disabled")

logger = logging.getLogger(__name__)

class ExportFormat:
    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"
    JSON = "json"
    MD = "md"
    HTML = "html"

class ExportSystem:
    """Comprehensive export system for all content types"""
    
    def __init__(self):
        self.supported_formats = [
            ExportFormat.TXT,
            ExportFormat.JSON,
            ExportFormat.MD,
            ExportFormat.HTML
        ]
        
        if PDF_AVAILABLE:
            self.supported_formats.append(ExportFormat.PDF)
        
        if DOCX_AVAILABLE:
            self.supported_formats.append(ExportFormat.DOCX)

    def export_poem(self, poem_data: Dict[str, Any], format: str) -> bytes:
        """Export poem in specified format"""
        
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}")
        
        title = poem_data.get('title', 'Untitled Poem')
        content = poem_data.get('content', '')
        poem_type = poem_data.get('type', 'poem')
        metadata = poem_data.get('metadata', {})
        
        if format == ExportFormat.TXT:
            return self._export_poem_txt(title, content, poem_type, metadata)
        elif format == ExportFormat.PDF:
            return self._export_poem_pdf(title, content, poem_type, metadata)
        elif format == ExportFormat.DOCX:
            return self._export_poem_docx(title, content, poem_type, metadata)
        elif format == ExportFormat.JSON:
            return self._export_poem_json(poem_data)
        elif format == ExportFormat.MD:
            return self._export_poem_md(title, content, poem_type, metadata)
        elif format == ExportFormat.HTML:
            return self._export_poem_html(title, content, poem_type, metadata)

    def export_story(self, story_data: Dict[str, Any], format: str) -> bytes:
        """Export story in specified format"""
        
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}")
        
        title = story_data.get('title', 'Untitled Story')
        content = story_data.get('content', '')
        outline = story_data.get('outline')
        metadata = story_data.get('metadata', {})
        
        if format == ExportFormat.TXT:
            return self._export_story_txt(title, content, outline, metadata)
        elif format == ExportFormat.PDF:
            return self._export_story_pdf(title, content, outline, metadata)
        elif format == ExportFormat.DOCX:
            return self._export_story_docx(title, content, outline, metadata)
        elif format == ExportFormat.JSON:
            return self._export_story_json(story_data)
        elif format == ExportFormat.MD:
            return self._export_story_md(title, content, outline, metadata)
        elif format == ExportFormat.HTML:
            return self._export_story_html(title, content, outline, metadata)

    def export_writing(self, writing_data: Dict[str, Any], format: str) -> bytes:
        """Export writing in specified format"""
        
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}")
        
        title = writing_data.get('title', 'Untitled Writing')
        content = writing_data.get('content', '')
        writing_type = writing_data.get('writing_type', 'document')
        metadata = writing_data.get('metadata', {})
        
        if format == ExportFormat.TXT:
            return self._export_writing_txt(title, content, writing_type, metadata)
        elif format == ExportFormat.PDF:
            return self._export_writing_pdf(title, content, writing_type, metadata)
        elif format == ExportFormat.DOCX:
            return self._export_writing_docx(title, content, writing_type, metadata)
        elif format == ExportFormat.JSON:
            return self._export_writing_json(writing_data)
        elif format == ExportFormat.MD:
            return self._export_writing_md(title, content, writing_type, metadata)
        elif format == ExportFormat.HTML:
            return self._export_writing_html(title, content, writing_type, metadata)

    def export_collection(self, items: List[Dict[str, Any]], format: str, 
                         collection_name: str = "Collection") -> bytes:
        """Export multiple items as a collection"""
        
        if format == ExportFormat.JSON:
            collection_data = {
                'collection_name': collection_name,
                'exported_at': datetime.now().isoformat(),
                'items': items,
                'count': len(items)
            }
            return json.dumps(collection_data, indent=2).encode('utf-8')
        
        elif format == ExportFormat.TXT:
            content_parts = [
                f"{collection_name}\n",
                "=" * len(collection_name) + "\n",
                f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n",
                f"Items: {len(items)}\n\n"
            ]
            
            for i, item in enumerate(items):
                content_parts.append(f"Item {i + 1}: {item.get('title', 'Untitled')}\n")
                content_parts.append("-" * 50 + "\n")
                content_parts.append(item.get('content', '') + "\n\n")
            
            return ''.join(content_parts).encode('utf-8')
        
        else:
            # For other formats, create a ZIP file with individual exports
            return self._export_collection_zip(items, format, collection_name)

    # TXT Export Methods
    def _export_poem_txt(self, title: str, content: str, poem_type: str, metadata: Dict) -> bytes:
        """Export poem as plain text"""
        lines = [
            title,
            "=" * len(title),
            f"Type: {poem_type.replace('_', ' ').title()}",
            f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            content,
            ""
        ]
        
        if metadata.get('syllable_counts'):
            lines.append("Syllable Pattern: " + " - ".join(str(x) for x in metadata['syllable_counts']))
        
        if metadata.get('rhyme_scheme'):
            lines.append(f"Rhyme Scheme: {metadata['rhyme_scheme']}")
        
        return '\n'.join(lines).encode('utf-8')

    def _export_story_txt(self, title: str, content: str, outline: Optional[Dict], metadata: Dict) -> bytes:
        """Export story as plain text"""
        lines = [
            title,
            "=" * len(title),
            f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ""
        ]
        
        if outline:
            lines.extend([
                "Story Details:",
                f"Genre: {outline.get('genre', 'Unknown').replace('_', ' ').title()}",
                f"Length: {outline.get('length', 'Unknown').replace('_', ' ').title()}",
                f"Word Count Target: {outline.get('word_count_target', 'Unknown')}",
                ""
            ])
        
        lines.extend([content, ""])
        
        return '\n'.join(lines).encode('utf-8')

    def _export_writing_txt(self, title: str, content: str, writing_type: str, metadata: Dict) -> bytes:
        """Export writing as plain text"""
        lines = [
            title,
            "=" * len(title),
            f"Type: {writing_type.replace('_', ' ').title()}",
            f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            content,
            ""
        ]
        
        return '\n'.join(lines).encode('utf-8')

    # JSON Export Methods
    def _export_poem_json(self, poem_data: Dict[str, Any]) -> bytes:
        """Export poem as JSON"""
        export_data = {
            **poem_data,
            'exported_at': datetime.now().isoformat(),
            'export_format': 'json'
        }
        return json.dumps(export_data, indent=2).encode('utf-8')

    def _export_story_json(self, story_data: Dict[str, Any]) -> bytes:
        """Export story as JSON"""
        export_data = {
            **story_data,
            'exported_at': datetime.now().isoformat(),
            'export_format': 'json'
        }
        return json.dumps(export_data, indent=2).encode('utf-8')

    def _export_writing_json(self, writing_data: Dict[str, Any]) -> bytes:
        """Export writing as JSON"""
        export_data = {
            **writing_data,
            'exported_at': datetime.now().isoformat(),
            'export_format': 'json'
        }
        return json.dumps(export_data, indent=2).encode('utf-8')

    # Markdown Export Methods
    def _export_poem_md(self, title: str, content: str, poem_type: str, metadata: Dict) -> bytes:
        """Export poem as Markdown"""
        lines = [
            f"# {title}",
            "",
            f"**Type:** {poem_type.replace('_', ' ').title()}  ",
            f"**Created:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ""
        ]
        
        if metadata.get('syllable_counts'):
            lines.append(f"**Syllable Pattern:** {' - '.join(str(x) for x in metadata['syllable_counts'])}  ")
        
        if metadata.get('rhyme_scheme'):
            lines.append(f"**Rhyme Scheme:** {metadata['rhyme_scheme']}  ")
        
        lines.extend(["", "---", "", content, ""])
        
        return '\n'.join(lines).encode('utf-8')

    def _export_story_md(self, title: str, content: str, outline: Optional[Dict], metadata: Dict) -> bytes:
        """Export story as Markdown"""
        lines = [
            f"# {title}",
            ""
        ]
        
        if outline:
            lines.extend([
                "## Story Details",
                "",
                f"- **Genre:** {outline.get('genre', 'Unknown').replace('_', ' ').title()}",
                f"- **Length:** {outline.get('length', 'Unknown').replace('_', ' ').title()}",
                f"- **Word Count Target:** {outline.get('word_count_target', 'Unknown')}",
                f"- **Created:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                "",
                "---",
                ""
            ])
        
        lines.extend([content, ""])
        
        return '\n'.join(lines).encode('utf-8')

    def _export_writing_md(self, title: str, content: str, writing_type: str, metadata: Dict) -> bytes:
        """Export writing as Markdown"""
        lines = [
            f"# {title}",
            "",
            f"**Type:** {writing_type.replace('_', ' ').title()}  ",
            f"**Created:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "---",
            "",
            content,
            ""
        ]
        
        return '\n'.join(lines).encode('utf-8')

    # HTML Export Methods
    def _export_poem_html(self, title: str, content: str, poem_type: str, metadata: Dict) -> bytes:
        """Export poem as HTML"""
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .header {{ border-bottom: 2px solid #333; margin-bottom: 30px; padding-bottom: 10px; }}
        .metadata {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 30px; }}
        .poem-content {{ white-space: pre-line; font-size: 1.1em; text-align: center; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
    </div>
    <div class="metadata">
        <p><strong>Type:</strong> {poem_type.replace('_', ' ').title()}</p>
        <p><strong>Created:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        {f'<p><strong>Syllable Pattern:</strong> {" - ".join(str(x) for x in metadata["syllable_counts"])}</p>' if metadata.get('syllable_counts') else ''}
        {f'<p><strong>Rhyme Scheme:</strong> {metadata["rhyme_scheme"]}</p>' if metadata.get('rhyme_scheme') else ''}
    </div>
    <div class="poem-content">
        {content}
    </div>
</body>
</html>
        """.strip()
        
        return html_content.encode('utf-8')

    def _export_story_html(self, title: str, content: str, outline: Optional[Dict], metadata: Dict) -> bytes:
        """Export story as HTML"""
        outline_section = ""
        if outline:
            outline_section = f"""
    <div class="metadata">
        <h2>Story Details</h2>
        <p><strong>Genre:</strong> {outline.get('genre', 'Unknown').replace('_', ' ').title()}</p>
        <p><strong>Length:</strong> {outline.get('length', 'Unknown').replace('_', ' ').title()}</p>
        <p><strong>Word Count Target:</strong> {outline.get('word_count_target', 'Unknown')}</p>
        <p><strong>Created:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
            """
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .header {{ border-bottom: 2px solid #333; margin-bottom: 30px; padding-bottom: 10px; }}
        .metadata {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 30px; }}
        .story-content {{ white-space: pre-line; font-size: 1.1em; text-align: justify; }}
        p {{ margin-bottom: 1em; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
    </div>
    {outline_section}
    <div class="story-content">
        {content.replace(chr(10), '</p><p>') if content else ''}
    </div>
</body>
</html>
        """.strip()
        
        return html_content.encode('utf-8')

    def _export_writing_html(self, title: str, content: str, writing_type: str, metadata: Dict) -> bytes:
        """Export writing as HTML"""
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        .header {{ border-bottom: 2px solid #333; margin-bottom: 30px; padding-bottom: 10px; }}
        .metadata {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 30px; }}
        .writing-content {{ white-space: pre-line; font-size: 1.1em; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
    </div>
    <div class="metadata">
        <p><strong>Type:</strong> {writing_type.replace('_', ' ').title()}</p>
        <p><strong>Created:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    <div class="writing-content">
        {content}
    </div>
</body>
</html>
        """.strip()
        
        return html_content.encode('utf-8')

    # PDF Export Methods (if available)
    def _export_poem_pdf(self, title: str, content: str, poem_type: str, metadata: Dict) -> bytes:
        """Export poem as PDF"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF export not available - reportlab not installed")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            alignment=1,  # Center
            spaceAfter=20
        )
        story.append(Paragraph(title, title_style))
        
        # Metadata
        metadata_text = f"Type: {poem_type.replace('_', ' ').title()}<br/>"
        metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        if metadata.get('syllable_counts'):
            metadata_text += f"<br/>Syllable Pattern: {' - '.join(str(x) for x in metadata['syllable_counts'])}"
        
        if metadata.get('rhyme_scheme'):
            metadata_text += f"<br/>Rhyme Scheme: {metadata['rhyme_scheme']}"
        
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Content
        poem_style = ParagraphStyle(
            'PoemStyle',
            parent=styles['Normal'],
            fontSize=14,
            alignment=1,  # Center
            leading=20
        )
        
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, poem_style))
            else:
                story.append(Spacer(1, 10))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_story_pdf(self, title: str, content: str, outline: Optional[Dict], metadata: Dict) -> bytes:
        """Export story as PDF"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF export not available - reportlab not installed")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 20))
        
        # Metadata
        if outline:
            metadata_text = f"Genre: {outline.get('genre', 'Unknown').replace('_', ' ').title()}<br/>"
            metadata_text += f"Length: {outline.get('length', 'Unknown').replace('_', ' ').title()}<br/>"
            metadata_text += f"Target Word Count: {outline.get('word_count_target', 'Unknown')}<br/>"
            metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            story.append(Paragraph(metadata_text, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_writing_pdf(self, title: str, content: str, writing_type: str, metadata: Dict) -> bytes:
        """Export writing as PDF"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF export not available - reportlab not installed")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 20))
        
        # Metadata
        metadata_text = f"Type: {writing_type.replace('_', ' ').title()}<br/>"
        metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    # DOCX Export Methods (if available)
    def _export_poem_docx(self, title: str, content: str, poem_type: str, metadata: Dict) -> bytes:
        """Export poem as DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX export not available - python-docx not installed")
        
        doc = Document()
        
        # Title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        metadata_text = f"Type: {poem_type.replace('_', ' ').title()}\n"
        metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        if metadata.get('syllable_counts'):
            metadata_text += f"\nSyllable Pattern: {' - '.join(str(x) for x in metadata['syllable_counts'])}"
        
        if metadata.get('rhyme_scheme'):
            metadata_text += f"\nRhyme Scheme: {metadata['rhyme_scheme']}"
        
        doc.add_paragraph(metadata_text)
        doc.add_page_break()
        
        # Content
        for line in content.split('\n'):
            p = doc.add_paragraph(line if line.strip() else '')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_story_docx(self, title: str, content: str, outline: Optional[Dict], metadata: Dict) -> bytes:
        """Export story as DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX export not available - python-docx not installed")
        
        doc = Document()
        
        # Title
        doc.add_heading(title, level=1)
        
        # Metadata
        if outline:
            doc.add_heading('Story Details', level=2)
            metadata_text = f"Genre: {outline.get('genre', 'Unknown').replace('_', ' ').title()}\n"
            metadata_text += f"Length: {outline.get('length', 'Unknown').replace('_', ' ').title()}\n"
            metadata_text += f"Target Word Count: {outline.get('word_count_target', 'Unknown')}\n"
            metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            doc.add_paragraph(metadata_text)
            doc.add_page_break()
        
        # Content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_writing_docx(self, title: str, content: str, writing_type: str, metadata: Dict) -> bytes:
        """Export writing as DOCX"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX export not available - python-docx not installed")
        
        doc = Document()
        
        # Title
        doc.add_heading(title, level=1)
        
        # Metadata
        metadata_text = f"Type: {writing_type.replace('_', ' ').title()}\n"
        metadata_text += f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        doc.add_paragraph(metadata_text)
        doc.add_page_break()
        
        # Content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _export_collection_zip(self, items: List[Dict[str, Any]], format: str, collection_name: str) -> bytes:
        """Export collection as ZIP file"""
        buffer = io.BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, item in enumerate(items):
                item_type = item.get('type', 'item')
                title = item.get('title', f'item_{i+1}')
                filename = f"{i+1:02d}_{title.replace(' ', '_')}.{format}"
                
                try:
                    if item_type == 'poem':
                        content = self.export_poem(item, format)
                    elif item_type == 'story':
                        content = self.export_story(item, format)
                    else:
                        content = self.export_writing(item, format)
                    
                    zip_file.writestr(filename, content)
                except Exception as e:
                    logger.error(f"Failed to export item {i+1}: {e}")
                    # Add error file
                    error_content = f"Failed to export: {str(e)}"
                    zip_file.writestr(f"{i+1:02d}_ERROR.txt", error_content.encode('utf-8'))
            
            # Add collection info
            info = {
                'collection_name': collection_name,
                'exported_at': datetime.now().isoformat(),
                'format': format,
                'item_count': len(items)
            }
            zip_file.writestr('collection_info.json', json.dumps(info, indent=2).encode('utf-8'))
        
        buffer.seek(0)
        return buffer.getvalue()

    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        return self.supported_formats.copy()

def create_export_system() -> ExportSystem:
    """Factory function to create export system"""
    return ExportSystem()